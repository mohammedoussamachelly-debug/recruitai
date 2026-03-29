"""
cv_parser.py — PDF and Word CV text extraction module.
Uses pdfplumber for PDFs and python-docx for .docx files.
"""

import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file) -> str:
    """
    Extract all text from a PDF file.
    Args:
        file: A file-like object (BytesIO or UploadedFile) or a file path string.
    Returns:
        Extracted text as a single string.
    """
    text_parts = []
    if isinstance(file, str):
        pdf = pdfplumber.open(file)
    else:
        pdf = pdfplumber.open(io.BytesIO(file.read() if hasattr(file, 'read') else file))

    with pdf as p:
        for page in p.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_text_from_docx(file) -> str:
    """
    Extract all text from a Word (.docx) file.
    Args:
        file: A file-like object (BytesIO or UploadedFile) or a file path string.
    Returns:
        Extracted text as a single string.
    """
    if isinstance(file, str):
        doc = Document(file)
    else:
        content = file.read() if hasattr(file, 'read') else file
        doc = Document(io.BytesIO(content))

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_cv(file, filename: str) -> str:
    """
    Dispatcher: detect the file type by extension and extract text.
    Args:
        file: A file-like object or bytes.
        filename: Original filename (used to detect extension).
    Returns:
        Extracted text string.
    Raises:
        ValueError: If the file type is not supported.
    """
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF or DOCX file.")
